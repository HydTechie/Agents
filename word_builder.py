from docx import Document
from docx.shared import Inches
import os

def build_doc(results, output_path):
    doc = Document()
    for pdf_name, flows in results.items():
        doc.add_heading(pdf_name, 0)
        for flow in flows:
            doc.add_heading(flow["name"], 1)
            if os.path.exists(flow["image"]):
                doc.add_picture(flow["image"], width=Inches(6))
    doc.save(output_path)
